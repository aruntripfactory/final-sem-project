# utils/export_tools.py
"""
Export Tools for Chat History and Analysis Results
Supports PDF, Word, and CSV export formats
"""

import pandas as pd
from datetime import datetime
from typing import List, Dict
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io


class ChatExporter:
    """Export chat history and analysis results"""
    
    def __init__(self):
        self.export_metadata = {
            'export_time': datetime.now(),
            'tool_name': 'AI Research Literature Assistant'
        }
    
    def export_to_pdf(
        self,
        chat_history: List[Dict],
        filename: str = "chat_export.pdf",
        include_metadata: bool = False
    ) -> str:
        """Export simplified chat history to PDF (Q&A only)"""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        pdf.set_font("Arial", '', 11)
        
        for i, entry in enumerate(chat_history, 1):
            role = entry.get('role', 'unknown').capitalize()
            content = entry.get('content', 'N/A')
            
            # Format as: [Role] Content
            text = f"[{role}] {content}"
            pdf.multi_cell(0, 7, txt=self._clean_text(text))
            pdf.ln(5)
            
            if pdf.get_y() > 270:
                pdf.add_page()
        
        pdf.output(filename)
        return filename
    
    def export_to_word(
        self,
        chat_history: List[Dict],
        filename: str = "chat_export.docx",
        include_metadata: bool = False
    ) -> str:
        """Export simplified chat history to Word (Q&A only)"""
        doc = Document()
        
        for entry in chat_history:
            role = entry.get('role', 'unknown').capitalize()
            content = entry.get('content', 'N/A')
            
            p = doc.add_paragraph()
            run = p.add_run(f"[{role}] ")
            run.bold = True
            p.add_run(content)
            
        doc.save(filename)
        return filename
    
    def export_to_csv(
        self,
        chat_history: List[Dict],
        filename: str = "chat_export.csv"
    ) -> str:
        """Export simplified chat history to CSV (Q&A only)"""
        data = []
        for entry in chat_history:
            data.append({
                'Role': entry.get('role', 'unknown').capitalize(),
                'Content': entry.get('content', 'N/A')
            })
        
        df = pd.DataFrame(data)
        df.to_csv(filename, index=False, encoding='utf-8')
        return filename
    
    def export_comparison_table(
        self,
        comparison_df: pd.DataFrame,
        filename: str = "comparison.xlsx"
    ) -> str:
        """
        Export comparison table to Excel
        
        Args:
            comparison_df: DataFrame with comparison data
            filename: Output filename
            
        Returns:
            Path to generated Excel file
        """
        with pd.ExcelWriter(filename, engine='openpyxl') as writer:
            comparison_df.to_excel(writer, sheet_name='Comparison', index=False)
            
            # Auto-adjust column widths
            worksheet = writer.sheets['Comparison']
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        return filename
    
    def export_to_json(
        self,
        chat_history: List[Dict],
        filename: str = "chat_export.json"
    ) -> str:
        """Export simplified chat history to JSON"""
        import json
        data = []
        for entry in chat_history:
            data.append({
                'role': entry.get('role', 'unknown'),
                'content': entry.get('content', '')
            })
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        return filename

    def export_to_txt(
        self,
        chat_history: List[Dict],
        filename: str = "chat_export.txt"
    ) -> str:
        """Export simplified chat history to TXT (Q&A only)"""
        with open(filename, 'w', encoding='utf-8') as f:
            for entry in chat_history:
                role = entry.get('role', 'unknown').upper()
                content = entry.get('content', '')
                f.write(f"[{role}] {content}\n\n")
        return filename

    def export_chat(
        self,
        chat_history: List[Dict],
        format_type: str = "txt",
        session_id: str = "current"
    ) -> str:
        """
        Unified export method
        
        Args:
            chat_history: List of chat messages
            format_type: One of 'json', 'csv', 'txt', 'pdf', 'docx'
            session_id: Session identifier for filename
            
        Returns:
            Path to the exported file
        """
        format_type = format_type.lower()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"chat_export_{session_id}_{timestamp}.{format_type}"
        
        if format_type == 'json':
            return self.export_to_json(chat_history, filename)
        elif format_type == 'csv':
            return self.export_to_csv(chat_history, filename)
        elif format_type == 'txt':
            return self.export_to_txt(chat_history, filename)
        elif format_type == 'pdf':
            return self.export_to_pdf(chat_history, filename)
        elif format_type == 'docx':
            return self.export_to_word(chat_history, filename)
        else:
            # Default to TXT
            return self.export_to_txt(chat_history, filename.replace(f'.{format_type}', '.txt'))

    def _clean_text(self, text: str) -> str:
        """Clean text for PDF export (remove problematic characters)"""
        if not text:
            return ""
        # Replace problematic characters
        replacements = {
            '\u2018': "'",  # Left single quote
            '\u2019': "'",  # Right single quote
            '\u201c': '"',  # Left double quote
            '\u201d': '"',  # Right double quote
            '\u2013': '-',  # En dash
            '\u2014': '-',  # Em dash
            '\u2022': '*',  # Bullet
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        # Remove non-ASCII characters
        text = text.encode('ascii', 'ignore').decode('ascii')
        return text


class SummarizationEngine:
    """Generate smart summaries at different levels"""
    
    SUMMARY_PROMPTS = {
        'executive': """Provide a concise 3-4 sentence executive summary covering:
        - Main contribution
        - Key methodology
        - Primary results
        
        Be specific and focus only on the most important information.""",
        
        'detailed': """Provide a comprehensive 1-page summary including:
        - Background and motivation
        - Methodology overview
        - Key results and findings
        - Main conclusions and limitations
        
        Use clear paragraphs and maintain academic tone.""",
        
        'sections': """Extract and summarize each major section:
        - Abstract (if present)
        - Introduction/Background
        - Methodology/Approach
        - Results/Findings
        - Conclusion/Discussion
        
        Provide 2-3 sentences for each section."""
    }
    
    @staticmethod
    def generate_summary(
        full_text: str,
        summary_type: str,
        generator_func,
        max_length: int = None
    ) -> str:
        """
        Generate summary of specified type
        
        Args:
            full_text: Full document text or chunks
            summary_type: Type of summary ('executive', 'detailed', 'sections')
            generator_func: LLM generation function
            max_length: Optional max length
            
        Returns:
            Generated summary
        """
        if summary_type not in SummarizationEngine.SUMMARY_PROMPTS:
            summary_type = 'executive'
        
        prompt = SummarizationEngine.SUMMARY_PROMPTS[summary_type]
        full_prompt = f"{prompt}\n\nDocument content:\n{full_text[:4000]}"  # Limit context
        
        try:
            summary = generator_func(full_prompt)
            
            if max_length and len(summary) > max_length:
                summary = summary[:max_length] + "..."
            
            return summary
        except Exception as e:
            return f"Error generating summary: {str(e)}"
    
    @staticmethod
    def extract_key_points(
        full_text: str,
        generator_func,
        num_points: int = 5
    ) -> List[str]:
        """
        Extract key bullet points from text
        
        Args:
            full_text: Full document text
            generator_func: LLM generation function
            num_points: Number of key points to extract
            
        Returns:
            List of key points
        """
        prompt = f"""Extract the {num_points} most important key points from this research paper.
        Format as a numbered list.
        
        Content:
        {full_text[:4000]}"""
        
        try:
            response = generator_func(prompt)
            # Parse into list
            lines = response.strip().split('\n')
            points = [line.strip() for line in lines if line.strip() and (line[0].isdigit() or line.startswith('-'))]
            return points[:num_points]
        except Exception as e:
            return [f"Error extracting key points: {str(e)}"]
